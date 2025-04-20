"""
Create a sample NVCA Series Seed SAFE template for demonstration purposes.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_series_seed_safe_template():
    """Create a sample Series Seed SAFE template with docxtpl fields."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header
    header = doc.add_heading('SAFE - SIMPLE AGREEMENT FOR FUTURE EQUITY', level=1)
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Company Info
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('THIS AGREEMENT is dated as of {{ effective_date }} and is between:').bold = True
    doc.add_paragraph()
    
    # Company name and details
    p = doc.add_paragraph()
    p.add_run('{{ company_name }}').bold = True
    p.add_run(', a Delaware corporation (the ')
    p.add_run('"Company"').bold = True
    p.add_run('),')
    
    doc.add_paragraph('and')
    
    # Investor details
    p = doc.add_paragraph()
    p.add_run('{{ investor_name }}').bold = True
    p.add_run(' (the ')
    p.add_run('"Investor"').bold = True
    p.add_run(').')
    
    doc.add_paragraph()
    
    # Introduction
    p = doc.add_paragraph()
    p.add_run('WHEREAS:').bold = True
    
    p = doc.add_paragraph('The Investor wishes to invest ${{ investment_amount }} in the Company on the terms set out in this agreement.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    p = doc.add_paragraph('The Company wishes to issue a SAFE (Simple Agreement for Future Equity) that converts into equity in the event of an equity financing, a liquidity event or a dissolution event on the terms set out in this agreement.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Agreement sections
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('IT IS AGREED as follows:').bold = True
    
    # Section 1: Definitions
    doc.add_paragraph()
    p = doc.add_paragraph('1. DEFINITIONS')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style = 'Heading 2'
    
    p = doc.add_paragraph('1.1 The following terms shall have the following meanings when used in this agreement:')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    p = doc.add_paragraph()
    p.add_run('"Valuation Cap" ').bold = True
    p.add_run('means ${{ valuation_cap }};')
    
    p = doc.add_paragraph()
    p.add_run('"Discount Rate" ').bold = True
    p.add_run('means {{ discount_rate }}%;')
    
    # Section 2: Investment Amount
    doc.add_paragraph()
    p = doc.add_paragraph('2. INVESTMENT AMOUNT')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style = 'Heading 2'
    
    p = doc.add_paragraph('2.1 The Investor hereby agrees to pay to the Company the sum of ${{ investment_amount }} (the "Investment Amount") by way of subscription for the SAFE, receipt of which the Company hereby acknowledges.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Section 3: Conversion
    doc.add_paragraph()
    p = doc.add_paragraph('3. CONVERSION')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style = 'Heading 2'
    
    p = doc.add_paragraph('3.1 In the event of an Equity Financing, the SAFE will automatically convert into the number of shares equal to the Investment Amount divided by the Conversion Price.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Section 4: Termination
    doc.add_paragraph()
    p = doc.add_paragraph('4. TERMINATION')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.style = 'Heading 2'
    
    p = doc.add_paragraph('4.1 This agreement shall terminate on the earlier of:')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    p = doc.add_paragraph('4.1.1 The issuance of shares to the Investor pursuant to the conversion of this SAFE; or')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    p = doc.add_paragraph('4.1.2 The payment, or setting aside for payment, of amounts due to the Investor pursuant to a Liquidity Event or a Dissolution Event.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Signature blocks
    doc.add_paragraph()
    doc.add_paragraph('IN WITNESS WHEREOF, the undersigned have caused this agreement to be duly executed and delivered.')
    doc.add_paragraph()
    
    p = doc.add_paragraph('COMPANY:')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph('{{ company_name }}')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph()
    p.add_run('By: ____________________________')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph()
    p.add_run('Name: {{ company_signatory_name }}')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph()
    p.add_run('Title: {{ company_signatory_title }}')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('INVESTOR:')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph('{{ investor_name }}')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph()
    p.add_run('By: ____________________________')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph()
    p.add_run('Name: {{ investor_signatory_name }}')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    p = doc.add_paragraph()
    p.add_run('Title: {{ investor_signatory_title }}')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Save the document
    os.makedirs('nvca', exist_ok=True)
    doc.save('nvca/SeriesSeed-SAFE-Template.docx')
    print("Template created successfully!")

if __name__ == "__main__":
    create_series_seed_safe_template()